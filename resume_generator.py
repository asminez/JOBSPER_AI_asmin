from datetime import datetime
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

class ResumeGenerator:
    def __init__(self):
        self.output_folder = 'output'
        os.makedirs(self.output_folder, exist_ok=True)
    
    def generate(self, resume_data, original_filename):
        """Generate templated resume"""
        doc = Document()
        
        # Setup document style
        self._setup_document_style(doc)
        
        # Add personal information
        self._add_personal_info(doc, resume_data.get('personal_info', {}))
        
        # Add professional summary
        if resume_data.get('summary'):
            self._add_section(doc, 'Professional Summary')
            self._add_paragraph(doc, resume_data['summary'])
        
        # Add work experience
        if resume_data.get('work_experience'):
            self._add_section(doc, 'Work Experience')
            for exp in resume_data['work_experience']:
                self._add_work_experience(doc, exp)
        
        # Add education
        if resume_data.get('education'):
            self._add_section(doc, 'Education')
            for edu in resume_data['education']:
                self._add_education(doc, edu)
        
        # Add skills
        if resume_data.get('skills'):
            self._add_section(doc, 'Skills')
            skills_text = ', '.join(resume_data['skills'])
            self._add_paragraph(doc, skills_text)
        
        # Add projects
        if resume_data.get('projects'):
            self._add_section(doc, 'Projects')
            for project in resume_data['projects']:
                self._add_project(doc, project)
        
        # Add certifications
        if resume_data.get('certifications'):
            self._add_section(doc, 'Certifications')
            for cert in resume_data['certifications']:
                self._add_certification(doc, cert)
        
        # Add languages
        if resume_data.get('languages'):
            self._add_section(doc, 'Languages')
            languages_text = ', '.join(resume_data['languages'])
            self._add_paragraph(doc, languages_text)
        
        # Add awards
        if resume_data.get('awards'):
            self._add_section(doc, 'Awards & Honors')
            for award in resume_data['awards']:
                self._add_award(doc, award)
        
        # Save file
        output_filename = f"resume_template_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        output_path = os.path.join(self.output_folder, output_filename)
        doc.save(output_path)
        
        return output_path
    
    def _setup_document_style(self, doc):
        """Setup document style"""
        # Set default font (supports English and Chinese)
        style = doc.styles['Normal']
        font = style.font
        # Use Chinese-compatible fonts, prefer Microsoft YaHei, fallback to SimSun
        try:
            font.name = 'Microsoft YaHei'  # Microsoft YaHei
        except:
            try:
                font.name = 'SimHei'  # SimHei (Bold)
            except:
                font.name = 'SimSun'  # SimSun (Song)
        font.size = Pt(11)
    
    def _add_personal_info(self, doc, personal_info):
        """Add personal information"""
        # Name
        if personal_info.get('name'):
            name_para = doc.add_paragraph()
            name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            name_run = name_para.add_run(personal_info['name'])
            name_run.font.size = Pt(18)
            name_run.font.bold = True
            name_run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Contact information
        contact_info = []
        if personal_info.get('email'):
            contact_info.append(personal_info['email'])
        if personal_info.get('phone'):
            contact_info.append(personal_info['phone'])
        if personal_info.get('address'):
            contact_info.append(personal_info['address'])
        if personal_info.get('linkedin'):
            contact_info.append(personal_info['linkedin'])
        if personal_info.get('github'):
            contact_info.append(personal_info['github'])
        if personal_info.get('website'):
            contact_info.append(personal_info['website'])
        
        if contact_info:
            contact_para = doc.add_paragraph()
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_run = contact_para.add_run(' | '.join(contact_info))
            contact_run.font.size = Pt(10)
        
        doc.add_paragraph()  # Empty line
    
    def _add_section(self, doc, title):
        """Add section title"""
        para = doc.add_paragraph()
        run = para.add_run(title.upper())
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(6)
    
    def _add_paragraph(self, doc, text):
        """Add regular paragraph"""
        if text:
            para = doc.add_paragraph(text)
            para.paragraph_format.space_after = Pt(6)
    
    def _add_work_experience(self, doc, exp):
        """Add work experience"""
        # Position and company
        header_text = []
        if exp.get('position'):
            header_text.append(exp['position'])
        if exp.get('company'):
            header_text.append(f"at {exp['company']}")
        
        if header_text:
            para = doc.add_paragraph()
            run = para.add_run(' | '.join(header_text))
            run.font.bold = True
            run.font.size = Pt(11)
        
        # Time period and location
        period_location = []
        if exp.get('period'):
            period_location.append(exp['period'])
        if exp.get('location'):
            period_location.append(exp['location'])
        
        if period_location:
            para = doc.add_paragraph()
            run = para.add_run(' | '.join(period_location))
            run.font.size = Pt(10)
            run.font.italic = True
        
        # Description
        if exp.get('description'):
            if isinstance(exp['description'], list):
                for desc in exp['description']:
                    if desc:
                        para = doc.add_paragraph(desc, style='List Bullet')
            else:
                self._add_paragraph(doc, exp['description'])
        
        doc.add_paragraph()  # Empty line
    
    def _add_education(self, doc, edu):
        """Add education background"""
        # Degree and major
        header_text = []
        if edu.get('degree'):
            header_text.append(edu['degree'])
        if edu.get('major'):
            header_text.append(f"in {edu['major']}")
        
        if header_text:
            para = doc.add_paragraph()
            run = para.add_run(' | '.join(header_text))
            run.font.bold = True
            run.font.size = Pt(11)
        
        # Institution
        if edu.get('institution'):
            para = doc.add_paragraph(edu['institution'])
            para.runs[0].font.size = Pt(10)
        
        # Time period and GPA
        details = []
        if edu.get('period'):
            details.append(edu['period'])
        if edu.get('gpa'):
            details.append(f"GPA: {edu['gpa']}")
        
        if details:
            para = doc.add_paragraph(' | '.join(details))
            para.runs[0].font.size = Pt(10)
            para.runs[0].font.italic = True
        
        doc.add_paragraph()  # Empty line
    
    def _add_project(self, doc, project):
        """Add project experience"""
        # Project name
        if project.get('name'):
            para = doc.add_paragraph()
            run = para.add_run(project['name'])
            run.font.bold = True
            run.font.size = Pt(11)
        
        # Description
        if project.get('description'):
            self._add_paragraph(doc, project['description'])
        
        # Technologies
        if project.get('technologies'):
            tech_text = f"Technologies: {', '.join(project['technologies'])}"
            para = doc.add_paragraph(tech_text)
            para.runs[0].font.size = Pt(10)
            para.runs[0].font.italic = True
        
        doc.add_paragraph()  # Empty line
    
    def _add_certification(self, doc, cert):
        """Add certification"""
        cert_text = []
        if cert.get('name'):
            cert_text.append(cert['name'])
        if cert.get('issuer'):
            cert_text.append(f"({cert['issuer']})")
        if cert.get('date'):
            cert_text.append(f"- {cert['date']}")
        
        if cert_text:
            para = doc.add_paragraph(' '.join(cert_text))
            para.paragraph_format.space_after = Pt(3)
    
    def _add_award(self, doc, award):
        """Add award"""
        award_text = []
        if award.get('name'):
            award_text.append(award['name'])
        if award.get('date'):
            award_text.append(f"({award['date']})")
        
        if award_text:
            para = doc.add_paragraph(' - '.join(award_text))
            para.paragraph_format.space_after = Pt(3)

